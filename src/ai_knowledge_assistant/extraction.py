"""Deterministic extraction from validated, workspace-owned documents only."""

from __future__ import annotations

import hashlib
import os
import re
from io import BytesIO
from pathlib import Path, PureWindowsPath
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from .models import (
    AcceptedDocument,
    DocumentType,
    ExtractedDocument,
    ExtractedSection,
    ExtractionError,
    ExtractionErrorCode,
    SourceLocator,
    SourceLocatorKind,
)
from .workspace import UploadWorkspace

EXTRACTION_VERSION = "v1"
MAX_DOCX_UNCOMPRESSED_SIZE_BYTES = 50 * 1024 * 1024
_MARKDOWN_HEADING = re.compile(r"^(#{1,6})[ \t]+(.+?)[ \t]*#*[ \t]*$")
_DOCX_HEADING_STYLE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)


def extract_document(
    workspace: UploadWorkspace, accepted_document: AcceptedDocument
) -> ExtractedDocument:
    """Extract text from one accepted file owned by ``workspace``.

    This boundary uses no network, AI, OCR, rendering, or external resource
    loading. It reads the accepted file once and verifies its exact-byte hash
    before passing bytes to the format-specific extractor.
    """
    content = read_accepted_content(workspace, accepted_document)
    if accepted_document.document_type is DocumentType.PDF:
        sections = _extract_pdf(accepted_document, content)
    elif accepted_document.document_type is DocumentType.DOCX:
        sections = _extract_docx(accepted_document, content)
    elif accepted_document.document_type is DocumentType.TEXT:
        sections = _extract_text(accepted_document, content)
    elif accepted_document.document_type is DocumentType.MARKDOWN:
        sections = _extract_markdown(accepted_document, content)
    else:
        raise ExtractionError(
            ExtractionErrorCode.UNSUPPORTED_DOCUMENT_STRUCTURE,
            "This accepted document type cannot be extracted by V1.",
        )

    if not any(section.text.strip() for section in sections):
        raise ExtractionError(
            ExtractionErrorCode.NO_EXTRACTABLE_TEXT,
            "The document contains no extractable text.",
        )
    return ExtractedDocument(
        extraction_version=EXTRACTION_VERSION,
        document_id=accepted_document.document_id,
        document_display_name=accepted_document.original_display_name,
        document_type=accepted_document.document_type,
        source_content_hash=accepted_document.content_hash,
        sections=tuple(sections),
    )


def read_accepted_content(
    workspace: UploadWorkspace, accepted_document: AcceptedDocument
) -> bytes:
    _validate_workspace(workspace)
    if accepted_document.run_id != workspace.run_id:
        raise ExtractionError(
            ExtractionErrorCode.EXTRACTION_FAILED,
            "The document does not belong to this workspace.",
        )
    if (
        Path(accepted_document.stored_filename).name
        != accepted_document.stored_filename
        or "\\" in accepted_document.stored_filename
        or PureWindowsPath(accepted_document.stored_filename).drive
    ):
        raise ExtractionError(
            ExtractionErrorCode.EXTRACTION_FAILED,
            "The stored document reference is invalid.",
        )
    directory_flags = os.O_RDONLY
    if hasattr(os, "O_DIRECTORY"):
        directory_flags |= os.O_DIRECTORY
    if hasattr(os, "O_NOFOLLOW"):
        directory_flags |= os.O_NOFOLLOW
    file_flags = os.O_RDONLY
    if hasattr(os, "O_NOFOLLOW"):
        file_flags |= os.O_NOFOLLOW
    try:
        directory_descriptor = os.open(workspace.uploads_dir, directory_flags)
        file_descriptor = os.open(
            accepted_document.stored_filename,
            file_flags,
            dir_fd=directory_descriptor,
        )
        with os.fdopen(file_descriptor, "rb") as stored_file:
            content = stored_file.read()
    except OSError as error:
        raise ExtractionError(
            ExtractionErrorCode.EXTRACTION_FAILED,
            "The stored document could not be read for extraction.",
        ) from error
    finally:
        if "directory_descriptor" in locals():
            os.close(directory_descriptor)
    if hashlib.sha256(content).hexdigest() != accepted_document.content_hash:
        raise ExtractionError(
            ExtractionErrorCode.EXTRACTION_FAILED,
            "The stored document no longer matches its accepted content.",
        )
    return content


def _validate_workspace(workspace: UploadWorkspace) -> None:
    if workspace.uploads_dir.parent != workspace.root:
        raise ExtractionError(
            ExtractionErrorCode.EXTRACTION_FAILED,
            "The workspace upload directory is invalid.",
        )
    if workspace.root.parent != workspace._temporary_root:
        raise ExtractionError(
            ExtractionErrorCode.EXTRACTION_FAILED,
            "The workspace is outside the controlled temporary root.",
        )
    if workspace.root.is_symlink() or workspace.uploads_dir.is_symlink():
        raise ExtractionError(
            ExtractionErrorCode.EXTRACTION_FAILED,
            "The workspace must not contain symbolic-link storage.",
        )
    if not workspace.uploads_dir.is_dir():
        raise ExtractionError(
            ExtractionErrorCode.EXTRACTION_FAILED,
            "The workspace upload directory is unavailable.",
        )


def _extract_text(document: AcceptedDocument, content: bytes) -> list[ExtractedSection]:
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError as error:
        raise ExtractionError(
            ExtractionErrorCode.EXTRACTION_FAILED,
            "The accepted text document could not be decoded as UTF-8.",
        ) from error
    normalized = _normalize_text(text)
    return [
        _section(
            document,
            SourceLocator(
                kind=SourceLocatorKind.TEXT_LINE_RANGE,
                line_start=1,
                line_end=max(1, len(normalized.splitlines())),
            ),
            normalized,
        )
    ]


def _extract_markdown(
    document: AcceptedDocument, content: bytes
) -> list[ExtractedSection]:
    try:
        source = content.decode("utf-8")
    except UnicodeDecodeError as error:
        raise ExtractionError(
            ExtractionErrorCode.EXTRACTION_FAILED,
            "The accepted Markdown document could not be decoded as UTF-8.",
        ) from error

    sections: list[ExtractedSection] = []
    current_label: str | None = None
    current_level: int | None = None
    current_lines: list[str] = []
    section_start = 1

    def flush(line_end: int) -> None:
        normalized = _normalize_text("\n".join(current_lines))
        if normalized.strip():
            sections.append(
                _section(
                    document,
                    SourceLocator(
                        kind=SourceLocatorKind.DOCUMENT_SECTION,
                        section_label=current_label,
                        heading_level=current_level,
                        line_start=section_start,
                        line_end=line_end,
                    ),
                    normalized,
                )
            )

    source_lines = _normalize_newlines(source).split("\n")
    for line_number, line in enumerate(source_lines, start=1):
        heading = _MARKDOWN_HEADING.match(line)
        if heading:
            flush(line_number - 1)
            current_label = heading.group(2).strip()
            current_level = len(heading.group(1))
            current_lines = [current_label]
            section_start = line_number
        else:
            current_lines.append(line)
    flush(max(1, len(source_lines)))
    return sections


def _extract_pdf(document: AcceptedDocument, content: bytes) -> list[ExtractedSection]:
    try:
        reader = PdfReader(BytesIO(content), strict=True)
    except (OSError, PdfReadError) as error:
        raise ExtractionError(
            ExtractionErrorCode.EXTRACTION_FAILED,
            "The PDF could not be read for text extraction.",
        ) from error

    try:
        pages = list(reader.pages)
    except PdfReadError as error:
        raise ExtractionError(
            ExtractionErrorCode.EXTRACTION_FAILED,
            "The PDF page structure could not be read for text extraction.",
        ) from error

    sections: list[ExtractedSection] = []
    for page_number, page in enumerate(pages, start=1):
        try:
            extracted_text = page.extract_text() or ""
        except Exception as error:
            raise ExtractionError(
                ExtractionErrorCode.PDF_PAGE_EXTRACTION_FAILED,
                f"Text could not be extracted from PDF page {page_number}.",
            ) from error
        sections.append(
            _section(
                document,
                SourceLocator(
                    kind=SourceLocatorKind.PDF_PAGE,
                    page_number=page_number,
                ),
                _normalize_text(extracted_text),
            )
        )
    return sections


def _extract_docx(document: AcceptedDocument, content: bytes) -> list[ExtractedSection]:
    _validate_docx_structure(content)
    try:
        word_document = Document(BytesIO(content))
    except (BadZipFile, KeyError, OSError, PackageNotFoundError) as error:
        raise ExtractionError(
            ExtractionErrorCode.DOCX_EXTRACTION_FAILED,
            "The DOCX document could not be opened for text extraction.",
        ) from error

    sections: list[ExtractedSection] = []
    current_label: str | None = None
    current_level: int | None = None
    current_paragraphs: list[str] = []
    section_start = 1

    def flush(paragraph_end: int) -> None:
        normalized = _normalize_text("\n\n".join(current_paragraphs))
        if normalized.strip():
            sections.append(
                _section(
                    document,
                    SourceLocator(
                        kind=SourceLocatorKind.DOCUMENT_SECTION,
                        section_label=current_label,
                        heading_level=current_level,
                        paragraph_start=section_start,
                        paragraph_end=paragraph_end,
                    ),
                    normalized,
                )
            )

    for paragraph_number, paragraph in enumerate(word_document.paragraphs, start=1):
        paragraph_text = paragraph.text
        heading_level = _heading_level(paragraph.style.name)
        if heading_level is not None:
            flush(paragraph_number - 1)
            current_label = paragraph_text
            current_level = heading_level
            current_paragraphs = [paragraph_text]
            section_start = paragraph_number
        else:
            current_paragraphs.append(paragraph_text)
    flush(max(1, len(word_document.paragraphs)))
    return sections


def _validate_docx_structure(content: bytes) -> None:
    try:
        with ZipFile(BytesIO(content)) as archive:
            members = archive.infolist()
    except BadZipFile as error:
        raise ExtractionError(
            ExtractionErrorCode.DOCX_EXTRACTION_FAILED,
            "The DOCX archive could not be inspected safely.",
        ) from error
    if sum(member.file_size for member in members) > MAX_DOCX_UNCOMPRESSED_SIZE_BYTES:
        raise ExtractionError(
            ExtractionErrorCode.UNSUPPORTED_DOCUMENT_STRUCTURE,
            "The DOCX document is too large after decompression for V1 extraction.",
        )
    if any(
        member.filename.startswith(("/", "\\"))
        or "\\" in member.filename
        or ".." in Path(member.filename).parts
        for member in members
    ):
        raise ExtractionError(
            ExtractionErrorCode.UNSUPPORTED_DOCUMENT_STRUCTURE,
            "The DOCX archive contains unsupported file paths.",
        )


def _heading_level(style_name: str) -> int | None:
    match = _DOCX_HEADING_STYLE.match(style_name)
    return int(match.group(1)) if match else None


def _normalize_newlines(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _normalize_text(text: str) -> str:
    """Apply conservative normalization without changing words or punctuation."""
    lines = [line.rstrip() for line in _normalize_newlines(text).split("\n")]
    normalized_lines: list[str] = []
    previous_blank = False
    for line in lines:
        is_blank = not line.strip()
        if is_blank and previous_blank:
            continue
        normalized_lines.append(line)
        previous_blank = is_blank
    return "\n".join(normalized_lines).rstrip()


def _section(
    document: AcceptedDocument, locator: SourceLocator, text: str
) -> ExtractedSection:
    section_key = "\x1f".join(
        (
            EXTRACTION_VERSION,
            document.document_id,
            locator.kind,
            str(locator.page_number),
            str(locator.section_label),
            str(locator.heading_level),
            str(locator.line_start),
            str(locator.line_end),
            str(locator.paragraph_start),
            str(locator.paragraph_end),
            text,
        )
    )
    return ExtractedSection(
        section_id=hashlib.sha256(section_key.encode("utf-8")).hexdigest(),
        document_id=document.document_id,
        document_display_name=document.original_display_name,
        document_type=document.document_type,
        source_locator=locator,
        text=text,
        content_hash=hashlib.sha256(text.encode("utf-8")).hexdigest(),
    )
