from __future__ import annotations

import hashlib
from dataclasses import replace
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from ai_knowledge_assistant.extraction import EXTRACTION_VERSION, extract_document
from ai_knowledge_assistant.models import ExtractionError, ExtractionErrorCode
from ai_knowledge_assistant.uploads import accept_upload
from ai_knowledge_assistant.workspace import UploadWorkspace


def _pdf_with_pages(*page_texts: str) -> bytes:
    writer = PdfWriter()
    for text in page_texts:
        page = writer.add_blank_page(width=612, height=792)
        if text:
            font = DictionaryObject(
                {
                    NameObject("/Type"): NameObject("/Font"),
                    NameObject("/Subtype"): NameObject("/Type1"),
                    NameObject("/BaseFont"): NameObject("/Helvetica"),
                }
            )
            page[NameObject("/Resources")] = DictionaryObject(
                {
                    NameObject("/Font"): DictionaryObject({NameObject("/F1"): font}),
                }
            )
            stream = DecodedStreamObject()
            stream.set_data(f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode())
            page[NameObject("/Contents")] = writer._add_object(stream)
    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def _docx_with_headings() -> bytes:
    document = Document()
    document.add_paragraph("Preamble first.")
    document.add_heading("Operations", level=1)
    document.add_paragraph("Open the register.")
    document.add_heading("Safety", level=2)
    document.add_paragraph("Escalate allergen concerns.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


@pytest.fixture
def workspace(tmp_path: Path) -> UploadWorkspace:
    created = UploadWorkspace.create(tmp_path)
    yield created
    created.cleanup()


def test_txt_extraction_normalizes_conservatively(workspace: UploadWorkspace) -> None:
    accepted = accept_upload(
        workspace, "notes.txt", b"First\r\n\r\n\r\nSecond  \rThird"
    )

    extracted = extract_document(workspace, accepted)

    assert extracted.extraction_version == EXTRACTION_VERSION
    assert extracted.sections[0].text == "First\n\nSecond\nThird"
    assert extracted.sections[0].source_locator.line_start == 1
    assert extracted.sections[0].source_locator.line_end == 4


def test_markdown_preserves_heading_boundaries_and_order(
    workspace: UploadWorkspace,
) -> None:
    accepted = accept_upload(
        workspace,
        "guide.md",
        (
            b"Intro text.\n\n# First\nFirst text.\n\n## Child\nChild text."
            b"\n\n# Last\nLast text."
        ),
    )

    extracted = extract_document(workspace, accepted)

    assert [section.source_locator.section_label for section in extracted.sections] == [
        None,
        "First",
        "Child",
        "Last",
    ]
    assert [section.source_locator.heading_level for section in extracted.sections] == [
        None,
        1,
        2,
        1,
    ]
    assert extracted.sections[2].text == "Child\nChild text."


def test_pdf_extraction_preserves_page_numbers_order_and_empty_pages(
    workspace: UploadWorkspace,
) -> None:
    accepted = accept_upload(
        workspace, "guide.pdf", _pdf_with_pages("First", "", "Third")
    )

    extracted = extract_document(workspace, accepted)

    assert [section.source_locator.page_number for section in extracted.sections] == [
        1,
        2,
        3,
    ]
    assert [section.text for section in extracted.sections] == ["First", "", "Third"]


def test_docx_extraction_preserves_headings_and_paragraph_order(
    workspace: UploadWorkspace,
) -> None:
    accepted = accept_upload(workspace, "guide.docx", _docx_with_headings())

    extracted = extract_document(workspace, accepted)

    assert [section.source_locator.section_label for section in extracted.sections] == [
        None,
        "Operations",
        "Safety",
    ]
    assert [section.text for section in extracted.sections] == [
        "Preamble first.",
        "Operations\n\nOpen the register.",
        "Safety\n\nEscalate allergen concerns.",
    ]
    assert extracted.sections[1].source_locator.paragraph_start == 2
    assert extracted.sections[1].source_locator.paragraph_end == 3


def test_same_accepted_file_and_version_produce_identical_output(
    workspace: UploadWorkspace,
) -> None:
    accepted = accept_upload(workspace, "policy.md", b"# Policy\n\nUse written steps.")

    assert extract_document(workspace, accepted) == extract_document(
        workspace, accepted
    )


def test_extracted_metadata_does_not_include_absolute_paths(
    workspace: UploadWorkspace,
) -> None:
    accepted = accept_upload(workspace, "policy.txt", b"Only controlled content.")

    extracted = extract_document(workspace, accepted)

    assert str(workspace.root) not in repr(extracted)
    assert str(workspace.uploads_dir) not in repr(extracted)


def test_empty_text_and_image_only_pdf_return_no_extractable_text(
    workspace: UploadWorkspace,
) -> None:
    empty_text = accept_upload(workspace, "empty.txt", b"\r\n\r\n")
    empty_pdf = accept_upload(workspace, "empty.pdf", _pdf_with_pages(""))

    for accepted in (empty_text, empty_pdf):
        with pytest.raises(ExtractionError) as error:
            extract_document(workspace, accepted)
        assert error.value.code is ExtractionErrorCode.NO_EXTRACTABLE_TEXT


def test_malformed_pdf_and_docx_return_stable_extraction_errors(
    workspace: UploadWorkspace,
) -> None:
    pdf = accept_upload(workspace, "valid.pdf", _pdf_with_pages("Valid"))
    malformed_pdf = b"%PDF-broken"
    (workspace.uploads_dir / pdf.stored_filename).write_bytes(malformed_pdf)
    pdf = replace(pdf, content_hash=hashlib.sha256(malformed_pdf).hexdigest())

    with pytest.raises(ExtractionError) as error:
        extract_document(workspace, pdf)
    assert error.value.code is ExtractionErrorCode.EXTRACTION_FAILED

    docx = accept_upload(workspace, "valid.docx", _docx_with_headings())
    malformed_docx = b"PK\x03\x04"
    (workspace.uploads_dir / docx.stored_filename).write_bytes(malformed_docx)
    docx = replace(docx, content_hash=hashlib.sha256(malformed_docx).hexdigest())

    with pytest.raises(ExtractionError) as error:
        extract_document(workspace, docx)
    assert error.value.code is ExtractionErrorCode.DOCX_EXTRACTION_FAILED


def test_extraction_rejects_a_document_from_another_workspace(tmp_path: Path) -> None:
    first = UploadWorkspace.create(tmp_path)
    second = UploadWorkspace.create(tmp_path)
    accepted = accept_upload(first, "policy.txt", b"Scoped content.")

    with pytest.raises(ExtractionError) as error:
        extract_document(second, accepted)

    assert error.value.code is ExtractionErrorCode.EXTRACTION_FAILED
    first.cleanup()
    second.cleanup()


def test_extraction_rejects_changed_or_symlinked_stored_content(
    workspace: UploadWorkspace,
    tmp_path: Path,
) -> None:
    accepted = accept_upload(workspace, "policy.txt", b"Accepted content.")
    stored_path = workspace.uploads_dir / accepted.stored_filename
    stored_path.write_bytes(b"Changed content.")

    with pytest.raises(ExtractionError) as error:
        extract_document(workspace, accepted)
    assert error.value.code is ExtractionErrorCode.EXTRACTION_FAILED

    stored_path.unlink()
    outside_file = tmp_path / "outside.txt"
    outside_file.write_bytes(b"Accepted content.")
    stored_path.symlink_to(outside_file)

    with pytest.raises(ExtractionError) as error:
        extract_document(workspace, accepted)
    assert error.value.code is ExtractionErrorCode.EXTRACTION_FAILED
